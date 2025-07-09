import streamlit as st
import pandas as pd
import re
import io
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber
from pdf2image import convert_from_path
from PIL import Image

# ─── Configuración ─────────────────────────────────────────────────────────────
st.set_page_config(page_title="Generador de Informes TS", layout="wide")
st.title("📝 Generador de Informes Periciales")

BASE_DIR = Path(__file__).parent
BASE_TEMPLATE     = BASE_DIR / "PLANTILLA BASE - V3 generador.docx"
JURIDICO_TEMPLATE = BASE_DIR / "PLANTILLA BASE JURIDICO - V3 generador.docx"
POLIZA_XLSX       = BASE_DIR / "Modelos_de_poliza.xlsx"

IMG_PH = "{{IMG_CATASTRO}}"

# Regex de texto de encargo
REG_ENC = {
    "{{EXPEDIENTE}}": r"Expediente:\s*(\S+)",
    "{{FECHA_DE_OCURRENCIA}}": r"Fecha de Ocurrencia:\s*([0-9/]+)",
    "{{EFECTO}}": r"Efecto:\s*([0-9/]+)",
    "{{GARANTIA_AFECTADA}}": r"Garantia afectada:\s*(.+)",
    "{{FECHA_HORA_SERVICIO}}": r"<NI>([0-9]{2}-[0-9]{2}-[0-9]{4})",
    "{{ASEGURADO}}": r"Asegurado:\s*([^\r\n]+)",
    "{{TLF1}}": r"Tlf1\s*[:\-]?\s*([0-9]+)",
    "{{MODELO_CONDICIONES_GENERALES}}": r"MODELO CONDICIONES GENERALES:\s*([^\r\n]+)",
    "{{AGUA_CONTENIDO}}": r"AGUA CONTENIDO:\s*([0-9\.,]+)",
    "{{AGUA_CONTINENTE}}": r"AGUA CONTINENTE:\s*([0-9\.,]+)",
    "{{DIR_ENCARGO}}": r"Lugar:\s*([^\r\n]+)"
}

# Regex de catastro
REG_CAT = {
    "{{CP_CATASTRO}}": r"(\d{5})\s+",
    "{{LOCALIDAD_CATASTRO}}": r"\d{5}\s+([A-ZÁÉÍÓÚÜÑ\s]+)\[",
    "{{PROVINCIA_CATASTRO}}": r"\[([A-ZÁÉÍÓÚÜÑ]{2,})\]",
    "{{USO_PRINCIPAL_CATASTRAL}}": r"Uso principal:\s*([^\r\n]+?)(?:\s*Superficie|$)",
    "{{SUPERFICIE_CONSTRUIDA_CATASTRAL}}": r"Superficie construida:\s*([0-9.]+)",
    "{{ANO_CONSTRUCCION_CATASTRAL}}": r"Año construcción:\s*([0-9]{4})",
    "{{SUPERFICIE_ELEMENTOS_COMUNES}}": r"Elementos comunes[^0-9]*([0-9.]+)",
    "{{PARTICIPACION_INMUEBLE}}": r"Participación del inmueble:\s*([0-9\.,]+ ?%)"
}

# ─── Helpers ─────────────────────────────────────────────────────────────────────
def normaliza_modelo(m: str) -> str:
    c = m.upper().split()[0].split('ED.')[0]
    return re.sub(r'-GEN|-SXXI|-PCI|-CO|-TR','', c).rstrip('-')

def modelo_a_ramo(modelo: str) -> str:
    if not (modelo and POLIZA_XLSX.exists()): return ''
    df = pd.read_excel(POLIZA_XLSX, engine='openpyxl')
    clave = normaliza_modelo(modelo)
    row = df[df.iloc[:,0].str.startswith(clave, na=False)]
    return str(row.iloc[0,1]) if not row.empty else ''

def parse_encargo(text: str) -> dict:
    rep = {k:'' for k in REG_ENC}
    for k, pat in REG_ENC.items():
        m = re.search(pat, text, re.I)
        if m: rep[k] = m.group(1).strip()
    # Formatear fecha 2 dígitos
    f = rep.get("{{FECHA_DE_OCURRENCIA}}","")
    if f and len(f.split('/')[-1])==2:
        d,mo,y = f.split('/')
        rep["{{FECHA_DE_OCURRENCIA}}"] = f"{d}/{mo}/20{y}"
    rep["{{AGUA_CONTENIDO}}"]    = rep.get("{{AGUA_CONTENIDO}}","0")
    rep["{{AGUA_CONTINENTE}}"]   = rep.get("{{AGUA_CONTINENTE}}","0")
    rep["{{POLIZA_RAMO}}"]       = modelo_a_ramo(rep.get("{{MODELO_CONDICIONES_GENERALES}}",""))
    if rep.get("{{DIR_ENCARGO}}"): rep["{{DIR_CATASTRO}}"] = rep["{{DIR_ENCARGO}}"]
    return rep

def parse_catastro(file) -> dict:
    rep = {k:'' for k in REG_CAT}
    # Guardar temp
    with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.name).suffix) as tmp:
        tmp.write(file.read()); tmp.flush()
        path = tmp.name
    text = ""
    # Si es PDF extraer texto y página a imagen
    if path.lower().endswith(".pdf"):
        with pdfplumber.open(path) as pdf:
            text = pdf.pages[0].extract_text() or ""
        imgs = convert_from_path(path, first_page=1, last_page=1, dpi=200)
        img_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        imgs[0].save(img_tmp.name, "PNG")
        rep[IMG_PH] = img_tmp.name
    else:
        # imagen directa
        img = Image.open(path)
        img_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=Path(path).suffix)
        img.save(img_tmp.name)
        rep[IMG_PH] = img_tmp.name
    # Extraer datos catastrales
    for k, pat in REG_CAT.items():
        m = re.search(pat, text, re.I)
        if m: rep[k] = m.group(1).strip()
    return rep

def replace_runs_placeholder(runs, ph, val):
    # Inserta val en medio de runs reemplazando ph
    full = ''.join(r.text for r in runs)
    idx  = full.find(ph)
    if idx<0: return False
    start,end,pos = idx, idx+len(ph), 0
    for i,r in enumerate(runs):
        nxt = pos+len(r.text)
        if pos<=start<nxt: s_run,s_off = i, start-pos
        if pos<end<=nxt:  e_run,e_off = i, end-pos; break
        pos = nxt
    runs[s_run].text = runs[s_run].text[:s_off] + val + runs[e_run].text[e_off:]
    for j in range(s_run+1, e_run+1):
        runs[j].text = '' if j!=e_run else runs[j].text[e_off:]
    return True

def replace_all(doc: Document, rep: dict):
    # Párrafos
    for p in doc.paragraphs:
        if IMG_PH in p.text:
            for r in p.runs: r.clear()
            p.add_run().add_picture(rep[IMG_PH], width=Inches(6))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        for k,v in rep.items():
            if k in p.text:
                if not replace_runs_placeholder(p.runs, k, v):
                    p.text = p.text.replace(k,v)
    # Tablas
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if IMG_PH in p.text:
                        for r in p.runs: r.clear()
                        p.add_run().add_picture(rep[IMG_PH], width=Inches(6))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        continue
                    for k,v in rep.items():
                        if k in p.text:
                            if not replace_runs_placeholder(p.runs, k, v):
                                p.text = p.text.replace(k,v)

def add_photo_report(doc: Document, fotos):
    doc.add_page_break()
    doc.add_heading("Reportaje fotográfico", level=1)
    table = doc.add_table(rows=3, cols=2); table.autofit=True
    idx=0
    for r in range(3):
        for c in range(2):
            if idx>=len(fotos): break
            f = fotos[idx]
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=Path(f.name).suffix)
            tmp.write(f.read()); tmp.flush()
            cell = table.cell(r,c)
            cell.paragraphs[0].add_run().add_picture(tmp.name, width=Inches(2.5))
            idx+=1

# ─── Interfaz ───────────────────────────────────────────────────────────────────
with st.form("form"):
    st.subheader("1. Texto del encargo")
    txt_file  = st.file_uploader("", type=["txt"])
    txt_paste = st.text_area("… o pega aquí", height=150)
    st.subheader("2. Documento catastral (PDF o imagen)")
    cat_file = st.file_uploader("", type=["pdf","png","jpg","jpeg"])
    st.subheader("3. Fotografías (hasta 6 imágenes)")
    fotos = st.file_uploader("", type=["png","jpg","jpeg"], accept_multiple_files=True)
    st.subheader("4. Informe jurídico")
    is_jur = st.checkbox("Defensa jurídica")
    submitted = st.form_submit_button("Generar Informe")

if submitted:
    text = txt_file.read().decode("utf-8") if txt_file else txt_paste
    if not text:
        st.error("📌 Debes proporcionar el texto del encargo.")
        st.stop()
    rep = parse_encargo(text)
    if cat_file:
        rep.update(parse_catastro(cat_file))
    tpl = JURIDICO_TEMPLATE if is_jur or "ASIST.JURIDICA" in rep.get("{{GARANTIA_AFECTADA}}","") else BASE_TEMPLATE
    doc = Document(tpl)
    replace_all(doc, rep)
    if fotos:
        add_photo_report(doc, fotos)
    buf = io.BytesIO()
    filename = f"{rep.get('{{EXPEDIENTE}}','informe')}.docx"
    doc.save(buf); buf.seek(0)
    st.success("✅ Informe listo.")
    st.download_button("📥 Descargar informe", buf, file_name=filename,
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
